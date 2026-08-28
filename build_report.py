from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Examen_Muestreo_Resuelto.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
GRAY = "5E6872"
LIGHT_GRAY = "F2F4F7"
CODE_FILL = "F4F6F9"
RESULT_FILL = "E8F1F8"
BORDER = "C9D2DC"
WHITE = "FFFFFF"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_paragraph(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_paragraph_left_border(paragraph, color=BLUE, size="18", space="6"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)
    pbdr.append(left)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_run_font(run, size=9, color=GRAY)


def set_repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    tbl = table._tbl
    tblpr = tbl.tblPr

    for tag in ("w:tblW", "w:tblInd", "w:tblLayout", "w:tblCellMar", "w:tblBorders"):
        old = tblpr.find(qn(tag))
        if old is not None:
            tblpr.remove(old)

    tblw = OxmlElement("w:tblW")
    tblw.set(qn("w:w"), str(total))
    tblw.set(qn("w:type"), "dxa")
    tblpr.append(tblw)

    tblind = OxmlElement("w:tblInd")
    tblind.set(qn("w:w"), "120")
    tblind.set(qn("w:type"), "dxa")
    tblpr.append(tblind)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblpr.append(layout)

    margins = OxmlElement("w:tblCellMar")
    for side, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        item = OxmlElement(f"w:{side}")
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")
        margins.append(item)
    tblpr.append(margins)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER)
        borders.append(node)
    tblpr.append(borders)

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, widths_dxa, header=True):
    set_table_geometry(table, widths_dxa)
    if header:
        set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.0
                if c_idx > 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, size=9.5, bold=(r_idx == 0), color=INK if r_idx == 0 else None)


def set_up_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.10

    specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.line_spacing = 1.0

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = RGBColor.from_string(GRAY)
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)

    code = styles.add_style("CodigoR", 1)
    code.font.name = "Consolas"
    code.font.size = Pt(9)
    code.font.color.rgb = RGBColor.from_string(INK)
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(4)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.keep_together = True
    code.paragraph_format.keep_with_next = False

    answer = styles.add_style("Resultado", 1)
    answer.font.name = "Calibri"
    answer.font.size = Pt(11)
    answer._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    answer._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    answer.paragraph_format.left_indent = Inches(0.18)
    answer.paragraph_format.right_indent = Inches(0.12)
    answer.paragraph_format.space_before = Pt(0)
    answer.paragraph_format.space_after = Pt(6)
    answer.paragraph_format.line_spacing = 1.10
    answer.paragraph_format.keep_together = True


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def set_header_footer(section):
    for header in (section.header, section.even_page_header):
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        left = p.add_run("EIC401 | Muestreo aplicado en R")
        set_run_font(left, size=9, color=GRAY, bold=True)
        right = p.add_run("\tBase de tuberculosis | set.seed(2026)")
        set_run_font(right, size=9, color=GRAY)

    for footer in (section.footer, section.even_page_footer):
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        text = fp.add_run("EIC401 | Informe de resultados reproducibles")
        set_run_font(text, size=9, color=GRAY)


def add_metadata(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, size=11, color=INK, bold=True)
    value_run = p.add_run(value)
    set_run_font(value_run, size=11)


def add_code(doc, code):
    p = doc.add_paragraph(style="CodigoR")
    shade_paragraph(p, CODE_FILL)
    run = p.add_run("# Procedimiento en R\n" + code)
    set_run_font(run, name="Consolas", size=9, color=INK)
    return p


def add_result(doc, text, keep_with_next=False):
    p = doc.add_paragraph(style="Resultado")
    shade_paragraph(p, RESULT_FILL)
    set_paragraph_left_border(p)
    p.paragraph_format.keep_with_next = keep_with_next
    label = p.add_run("Resultado: ")
    set_run_font(label, size=11, color=INK, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_question(doc, number, prompt, code, result, code_before=2):
    h = doc.add_paragraph(style="Heading 2")
    h.add_run(f"{number}. {prompt}")
    code_paragraph = add_code(doc, code)
    code_paragraph.paragraph_format.space_before = Pt(code_before)
    add_result(doc, result)


def add_sampling_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Código de muestreo solicitado")
    set_run_font(run, size=11, color=INK, bold=True)
    add_code(doc, code)


def build_document():
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = False
    set_up_styles(doc)
    section = doc.sections[0]
    configure_section(section)

    # Primera página: memo masthead, conforme a standard_business_brief.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(28)
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(6)
    kr = kicker.add_run("INFORME DE RESULTADOS")
    set_run_font(kr, size=10, color=BLUE, bold=True)
    title = doc.add_paragraph("Muestreo aplicado en R", style="Title")
    subtitle = doc.add_paragraph(
        "Base de datos de tuberculosis | Muestreo aleatorio simple y estratificado",
        style="Subtitle",
    )
    add_metadata(doc, "Curso", "Software para Análisis e Interpretación de la Información para la Investigación Científica (EIC401)")
    add_metadata(doc, "Docente", "M.A. Lic. Jaime Andre Chocó Cedillos")
    add_metadata(doc, "Archivo analizado", "database_tb_limpia.xlsx")
    add_metadata(doc, "Población", "N = 3,104 pacientes")
    add_metadata(doc, "Semilla", "set.seed(2026)")
    add_metadata(doc, "Tamaño de cada muestra", "n = 300")
    add_metadata(doc, "Fecha", "27 de agosto de 2026")

    p = doc.add_paragraph(style="Resultado")
    shade_paragraph(p, RESULT_FILL)
    set_paragraph_left_border(p)
    r = p.add_run("Criterio de cálculo. ")
    set_run_font(r, bold=True, color=INK)
    r = p.add_run(
        "Las respuestas se calcularon únicamente con la muestra correspondiente. "
        "Las medias y medianas excluyen NA cuando el enunciado lo indica; los porcentajes de presencia usan como denominador el total de la muestra o del estrato, según corresponda."
    )
    set_run_font(r)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Reproducibilidad: ")
    set_run_font(r, bold=True, color=INK)
    r = p.add_run("R 4.4.2, readxl 1.4.5 y dplyr 1.1.4.")
    set_run_font(r, color=GRAY)

    doc.add_page_break()

    doc.add_paragraph("Ejercicio 1 - Muestreo aleatorio simple", style="Heading 1")
    add_sampling_block(
        doc,
        'library(readxl)\nlibrary(dplyr)\ndatos <- read_excel("database_tb_limpia.xlsx", sheet = "datos", guess_max = 5000)\ndatos <- as.data.frame(datos)\nset.seed(2026)\nn1 <- 300\nindices1 <- sample(1:nrow(datos), n1)\nmuestra1 <- datos[indices1, ]',
    )

    ex1 = [
        (
            "Tamaño de la muestra y porcentaje de la base",
            'nrow(muestra1)\n100 * nrow(muestra1) / nrow(datos)',
            "La muestra contiene 300 observaciones, equivalentes al 9.66% de los 3,104 registros de la base.",
        ),
        (
            "Distribución por sexo",
            'table(muestra1$sexo)\n100 * sum(muestra1$sexo == "Masculino") / nrow(muestra1)',
            "Se seleccionaron 190 hombres y 110 mujeres. Los hombres representan el 63.33% de la muestra.",
        ),
        (
            "Edad al diagnóstico: media, mediana y valores faltantes",
            'mean(muestra1$edad_al_diagnostico, na.rm = TRUE)\nmedian(muestra1$edad_al_diagnostico, na.rm = TRUE)\nsum(is.na(muestra1$edad_al_diagnostico))',
            "La edad media al diagnóstico es 36.35 años y la mediana es 34.00 años. La variable presenta 20 valores faltantes.",
        ),
        (
            "Porcentaje de pacientes fallecidos",
            '100 * sum(muestra1$pronostico == "Fallecido") / nrow(muestra1)',
            "Fallecieron 65 pacientes, lo que corresponde al 21.67% de la muestra.",
        ),
        (
            "Categoría más frecuente de edad_rangos y datos faltantes",
            'tabla_edad <- table(muestra1$edad_rangos)\nnames(which.max(tabla_edad))\nmax(tabla_edad)\nsum(is.na(muestra1$edad_rangos))',
            "La categoría más frecuente es “18 a 35 años”, con 121 pacientes. Hay 20 valores NA en edad_rangos (280 datos disponibles).",
        ),
        (
            "Porcentaje con coinfección",
            '100 * sum(muestra1$coinfeccion == "Sí") / nrow(muestra1)',
            "Presentan coinfección 158 pacientes, equivalentes al 52.67% de la muestra.",
        ),
        (
            "Diagnóstico de VIH positivo",
            'n_vih <- sum(muestra1$dx_vih == "Sí")\nc(n = n_vih, porcentaje = 100 * n_vih / nrow(muestra1))',
            "Tienen diagnóstico de VIH positivo 158 pacientes, equivalentes al 52.67%.",
        ),
        (
            "Clasificación según tipo_de_tb",
            'tabla_tipo <- table(muestra1$tipo_de_tb)\ntabla_tipo\nnames(which.max(tabla_tipo))',
            "Extrapulmonar: 141 casos; Pulmonar: 127; Profilaxis: 32. La categoría más frecuente es Extrapulmonar, con 141 casos.",
        ),
        (
            "Valores disponibles y media de cd4",
            'sum(!is.na(muestra1$cd4))\nmean(muestra1$cd4, na.rm = TRUE)',
            "Hay 84 valores disponibles de cd4. Su media es 140.99 células/ml.",
        ),
        (
            "Año con mayor número de pacientes",
            'tabla_anio <- table(muestra1$anio)\ntabla_anio\nnames(which.max(tabla_anio))',
            "El año con mayor representación es 2014, con 45 pacientes.",
        ),
        (
            "Pacientes con seguimiento Activo",
            'sum(!is.na(muestra1$seguimiento))\nsum(muestra1$seguimiento == "Activo", na.rm = TRUE)',
            "Los 300 pacientes tienen seguimiento registrado y 164 se encuentran en estado Activo.",
        ),
        (
            "Baciloscopia positiva y valores faltantes",
            'n_pos <- sum(muestra1$baciloscopia == "Sí", na.rm = TRUE)\n100 * n_pos / nrow(muestra1)\nsum(is.na(muestra1$baciloscopia))',
            "La baciloscopia es positiva en 230 pacientes, equivalentes al 76.67% de la muestra. Hay 21 valores NA.",
        ),
        (
            "Contacto con un caso de tuberculosis",
            'n_contacto <- sum(muestra1$conctacto == "Sí", na.rm = TRUE)\n100 * n_contacto / nrow(muestra1)',
            "Nueve pacientes reportan contacto con un caso de tuberculosis, equivalentes al 3.00% de la muestra. La variable se llama conctacto en la base.",
        ),
        (
            "Media y valores faltantes de meses_diagnostico_vih",
            'mean(muestra1$meses_diagnostico_vih, na.rm = TRUE)\nsum(is.na(muestra1$meses_diagnostico_vih))',
            "La media es 56.66 meses, calculada con los valores disponibles. La variable contiene 18 valores NA.",
        ),
    ]

    for idx, item in enumerate(ex1, start=1):
        add_question(doc, idx, *item)

    h = doc.add_paragraph(style="Heading 2")
    h.add_run("15. Tabla de frecuencias de cultivo e interpretación")
    add_code(
        doc,
        'tabla_cultivo <- table(muestra1$cultivo, useNA = "ifany")\nprop.table(tabla_cultivo)\n100 * prop.table(tabla_cultivo)',
    )
    add_result(
        doc,
        "El resultado Positivo es el más frecuente (155; 51.67%), seguido de 84 valores faltantes (28.00%) y 61 resultados Negativos (20.33%). En consecuencia, poco más de la mitad de la muestra tiene cultivo positivo; sin embargo, el 28.00% sin dato debe considerarse al interpretar la distribución.",
        keep_with_next=True,
    )
    table = doc.add_table(rows=1, cols=4)
    headers = ["Categoría", "Frecuencia", "Proporción", "Porcentaje"]
    for idx, value in enumerate(headers):
        table.rows[0].cells[idx].text = value
    rows = [
        ("Negativo", "61", "0.2033", "20.33%"),
        ("Positivo", "155", "0.5167", "51.67%"),
        ("NA", "84", "0.2800", "28.00%"),
        ("Total", "300", "1.0000", "100.00%"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = value
    style_table(table, [2880, 1800, 2160, 2520])
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)
    r = after.add_run("Fuente: cálculo propio con muestra1 (n = 300), set.seed(2026).")
    set_run_font(r, size=9, color=GRAY, italic=True)

    doc.add_page_break()
    doc.add_paragraph("Ejercicio 2 - Muestreo estratificado por sexo", style="Heading 1")
    add_sampling_block(
        doc,
        'datos <- read_excel("database_tb_limpia.xlsx", sheet = "datos", guess_max = 5000)\ndatos <- as.data.frame(datos)\n\nset.seed(2026)\nn_total <- 300\nprop_masc <- sum(datos$sexo == "Masculino") / nrow(datos)\nn_masc <- round(n_total * prop_masc)\nn_fem <- n_total - n_masc\n\ndatos_masc <- datos[datos$sexo == "Masculino", ]\ndatos_fem <- datos[datos$sexo == "Femenino", ]\n\nmuestra_masc <- datos_masc[sample(1:nrow(datos_masc), n_masc), ]\nmuestra_fem <- datos_fem[sample(1:nrow(datos_fem), n_fem), ]\nmuestra2 <- rbind(muestra_masc, muestra_fem)',
    )

    ex2 = [
        (
            "Composición de la muestra estratificada y verificación proporcional",
            'table(muestra2$sexo)\n100 * prop.table(table(muestra2$sexo))\n100 * prop.table(table(datos$sexo))',
            "La muestra contiene 189 hombres (63.00%) y 111 mujeres (37.00%). En la base completa hay 1,960 hombres (63.14%) y 1,144 mujeres (36.86%); por tanto, la asignación es consistente con la proporción poblacional, con la diferencia esperable por redondeo.",
        ),
        (
            "Porcentaje de fallecidos por sexo",
            'tab <- table(muestra2$sexo, muestra2$pronostico)\n100 * prop.table(tab, margin = 1)[, "Fallecido"]',
            "Falleció el 18.92% de las mujeres (21 de 111) y el 30.16% de los hombres (57 de 189).",
        ),
        (
            "Sexo con mayor porcentaje de fallecidos",
            'pct_fallecidos <- 100 * prop.table(tab, margin = 1)[, "Fallecido"]\nnames(which.max(pct_fallecidos))',
            "El porcentaje de fallecidos es mayor en los hombres: 30.16%, frente a 18.92% en las mujeres; la diferencia descriptiva es de 11.24 puntos porcentuales.",
        ),
        (
            "Edad promedio al diagnóstico por sexo",
            'aggregate(edad_al_diagnostico ~ sexo, muestra2,\n          FUN = function(x) mean(x, na.rm = TRUE))',
            "La edad promedio al diagnóstico es 35.80 años en hombres y 33.07 años en mujeres. Se excluyeron los valores NA del cálculo.",
        ),
        (
            "Mediana de la edad al diagnóstico por sexo",
            'aggregate(edad_al_diagnostico ~ sexo, muestra2,\n          FUN = function(x) median(x, na.rm = TRUE))',
            "La mediana es 33 años en hombres y 32 años en mujeres. Por lo tanto, la mediana es mayor en los hombres.",
        ),
        (
            "Diagnóstico de VIH positivo por sexo",
            'tab_vih <- table(muestra2$sexo, muestra2$dx_vih)\n100 * prop.table(tab_vih, margin = 1)[, "Sí"]',
            "El 54.50% de los hombres tiene diagnóstico de VIH positivo (103 de 189), frente al 45.95% de las mujeres (51 de 111).",
        ),
        (
            "Tipo de tuberculosis más frecuente por sexo",
            'tab_tipo <- table(muestra2$sexo, muestra2$tipo_de_tb)\napply(tab_tipo, 1, function(x) c(categoria = names(which.max(x)), n = max(x)))',
            "En ambos sexos predomina la tuberculosis Extrapulmonar: 86 casos en hombres y 58 casos en mujeres. En hombres siguen Pulmonar (81) y Profilaxis (22); en mujeres, Pulmonar (38) y Profilaxis (15).",
        ),
        (
            "Porcentaje con coinfección por sexo",
            'tab_coinf <- table(muestra2$sexo, muestra2$coinfeccion)\n100 * prop.table(tab_coinf, margin = 1)[, "Sí"]',
            "Presenta coinfección el 45.95% de las mujeres (51 de 111) y el 54.50% de los hombres (103 de 189).",
        ),
        (
            "Comparación del seguimiento Fallecido",
            'tab_seg <- table(muestra2$sexo, muestra2$seguimiento)\ntab_seg[, "Fallecido"]',
            "Hay 57 hombres y 21 mujeres con seguimiento Fallecido. El mayor número corresponde a los hombres.",
        ),
        (
            "Valores NA de edad_rangos por sexo",
            'tapply(is.na(muestra2$edad_rangos), muestra2$sexo, sum)',
            "No tienen dato en edad_rangos 12 mujeres y 6 hombres.",
        ),
        (
            "Baciloscopia positiva por sexo",
            'n_pos <- tapply(muestra2$baciloscopia == "Sí", muestra2$sexo,\n                function(x) sum(x, na.rm = TRUE))\n100 * n_pos / table(muestra2$sexo)',
            "La baciloscopia es positiva en el 71.43% de los hombres (135 de 189; 16 NA) y en el 70.27% de las mujeres (78 de 111; 14 NA). Los porcentajes usan el total de cada estrato como denominador.",
        ),
        (
            "Media de cd4 por sexo",
            'muestra2 %>% group_by(sexo) %>%\n  summarize(n = sum(!is.na(cd4)), media = mean(cd4, na.rm = TRUE))',
            "La media de cd4 es 138.95 células/ml en mujeres (20 datos disponibles) y 121.53 células/ml en hombres (58 disponibles). La media es mayor en las mujeres.",
        ),
        (
            "Rango de edad con mayor concentración por sexo",
            'tab_rango <- table(muestra2$sexo, muestra2$edad_rangos)\napply(tab_rango, 1, function(x) c(rango = names(which.max(x)), n = max(x)))',
            "En ambos sexos el rango más frecuente es 18 a 35 años: 39 mujeres y 84 hombres.",
        ),
        (
            "Casos de seguimiento Abandono por sexo",
            'tab_seg[, "Abandono"]',
            "Hay 1 caso de Abandono entre las mujeres y 6 entre los hombres.",
        ),
        (
            "Síntesis de diferencias entre hombres y mujeres",
            'sprintf("Fallecidos: H %.2f%%, M %.2f%%; abandono: H %d, M %d; CD4: H %.2f, M %.2f",\n        30.16, 18.92, 6, 1, 121.53, 138.95)',
            "Los hombres presentaron mayor mortalidad que las mujeres (30.16% frente a 18.92%), con 57 y 21 casos de seguimiento Fallecido, respectivamente. También hubo más abandonos entre los hombres (6 frente a 1) y una media de cd4 menor entre sus datos disponibles (121.53 frente a 138.95 células/ml). La positividad de la baciloscopia fue muy similar (71.43% en hombres y 70.27% en mujeres). En conjunto, la muestra muestra una evolución descriptivamente menos favorable en los hombres, sin que estos resultados demuestren causalidad.",
        ),
    ]

    for idx, item in enumerate(ex2, start=1):
        if idx == 11:
            doc.add_page_break()
        add_question(doc, idx, *item, code_before=8 if idx == 11 else 2)

    doc.add_paragraph("Nota final", style="Heading 1")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "Todas las cifras proceden de las muestras generadas con set.seed(2026). "
        "El archivo Examen_Muestreo_Resuelto.R contiene el código completo, comentado y ejecutable que reproduce estos resultados."
    )

    # Propiedades del documento.
    props = doc.core_properties
    props.title = "Muestreo aplicado en R - Base de datos de tuberculosis"
    props.subject = "Resultados de muestreo aleatorio simple y estratificado"
    props.author = "Estudiante EIC401"
    props.keywords = "R, muestreo, tuberculosis, set.seed(2026)"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
